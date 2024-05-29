# https://shravani1383-ai-travel-itinerary-generator-app-rrzl0w.streamlit.app/
import json
import logging
import math
import re
from datetime import timedelta
from io import BytesIO
import requests
import streamlit as st
import PIL.Image
from amadeus import Client, ResponseError
from docx import Document
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Mm
from docxtpl import DocxTemplate, InlineImage
from dotenv import set_key, load_dotenv, dotenv_values
from geopy.exc import GeocoderTimedOut
from geopy.geocoders import Nominatim
from openai import OpenAI
from spire.doc import *
from spire.doc.common import *
from streamlit.logger import get_logger
from table import generate_word_document
from docxcompose.composer import Composer
from docx import Document as Document_compose
import io

RAPID_API_HOST = "booking-com.p.rapidapi.com"
st.session_state['data_changed'] = False
input_dict = {}
st.set_page_config(
    page_title="AI Tour Itinerary Generator",  # Set your desired title here
    page_icon="images/favicon.ico",  # Set path to your favicon image (.ico format)
)
st.title("Tour Itinerary Generator")

# Check if the key exists
env_file_path = 'key.env'
env_vars = dotenv_values(env_file_path)

API_KEY = st.secrets['API_KEY']
RAPID_API_KEY = st.secrets["RAPID_API_KEY"]
AMADEUS_API_KEY = st.secrets["AMADEUS_API_KEY"]
AMADEUS_API_SECRET = st.secrets["AMADEUS_API_SECRET"]
PEXELS_API_KEY = st.secrets["PEXELS_API_KEY"]

col1, col2 = st.columns(2)

input_dict['dest'] = col1.text_input("Destination", key='dest',placeholder='ex. Himachal Pradesh')
input_dict['src'] = col1.text_input("Source City", key='src',placeholder='ex. Mumbai')
input_dict['genre'] = col1.text_input("Genre", key='genre',placeholder='ex. adventure, fun, religious')
input_dict['type_of_travelers'] = col1.text_input("Type of Travelers", key='type', placeholder='ex. family, friends')
input_dict['mode_of_travel'] = col1.text_input("Mode of Travel", key='mode', placeholder='ex. flight, bus, train')
input_dict['num_days'] = col2.number_input("Number of Days", key='num_days', min_value=0, max_value=None, value=0,
                                           step=1, format="%d")
input_dict['start_date'] = col2.date_input("Start Date", key='start_date')
# Create sub-columns within col2
col21, col22 = col2.columns(2)

input_dict['num_adults'] = int(
    col21.number_input("Number of Adults", key='num_adults', min_value=0, max_value=None, value=0, step=1, format="%d"))
input_dict['num_children'] = int(
    col22.number_input("Number of Children", key='num_children', min_value=0, max_value=None, value=0, step=1,
                       format="%d"))
input_dict['price_per_person'] = col2.number_input("Price Per Person", key='price_per_person', min_value=0)
input_dict['average_age'] = col2.number_input("Average age", key='average_age', min_value=0, max_value=None, value=0,
                                              step=1, format="%d")
# input_dict['food'] = 'non veg' if st.toggle('Include non-veg hotels') else 'veg'
special_note = st.text_area("Special Note(Optional)", key='special_note')

input_dict['num_tourists'] = input_dict['num_adults'] + input_dict['num_children']

client = OpenAI(api_key=API_KEY)

amadeus = Client(
    client_id='AMADEUS_API_KEY',
    client_secret='AMADEUS_API_SECRET'
)

function_descriptions = [
    {
        "name": "get_flight_hotel_info",
        "description": "Find the flights between cities and hotels within cities for residence.",
        "parameters": {
            "type": "object",
            "properties": {
                "loc_list": {
                    "type": "array",
                    "items": {
                        "type": "string"
                    },
                    "description": "The ordered list of names of cities in the tour. e.g. ['Mumbai', 'Paris']"
                },

                "date_list": {
                    "type": "array",
                    "items": {
                        "type": "string"
                    },
                    "description": "The ordered list of dates for arrival in the cities in YYYY-MM-DD format."
                },

                "iata_list": {
                    "type": "array",
                    "items": {
                        "type": "string"
                    },
                    "description": "The ordered list of IATA codes of cities in the tour which you have selected. e.g. ['BOM' for 'Mumbai', 'PAR' for 'Paris']"
                }

            },
            "required": ["loc_list", "date_list", "iata_list"],
        },
    }
]

# Set the logging level (DEBUG for most details)
logging.basicConfig(level=logging.DEBUG)

# Create a logger for your specific API usage
logger = logging.getLogger('booking_com_api')


def make_booking_com_api_call(url, headers, data):
    logger.debug("API Request:")
    logger.debug(f"URL: {url}")
    logger.debug(f"Headers: {headers}")
    logger.debug(f"Body: {data}")

    # Make your API call using requests library
    response = requests.get(url, headers=headers, json=data)

    logger.debug("API Response:")
    logger.debug(f"Status Code: {response.status_code}")
    logger.debug(f"Headers: {response.headers}")
    logger.debug(f"Response Body: {response.json()}")

    return response


def get_hotel_data(city, checkin_date, checkout_date, num_adults, num_children):
    city_dict = {}

    # Geocode city to get latitude and longitude
    def geocode_with_retry(location):
        geolocator = Nominatim(user_agent="my_geocoder")
        max_retries = 3

        for attempt in range(max_retries):
            try:
                geocoded_location = geolocator.geocode(location)
                return geocoded_location.raw
            except GeocoderTimedOut as e:
                print(f'Attempt {attempt + 1} failed: {e}')
                if attempt < max_retries - 1:
                    print(f'Retrying...')
                    continue

        print('Max retries exceeded. Unable to geocode the location.')
        return None

    location = geocode_with_retry(city)
    print('Location', location)

    if location:
        lat = location['lat']
        long = location['lon']

        # Define URL and query parameters for Booking.com API
        url = "https://booking-com.p.rapidapi.com/v1/hotels/search-by-coordinates"
        num_rooms = math.ceil((num_adults + num_children) / 3)
        querystring = {
            "locale": "en-gb",
            "room_number": num_rooms,
            "checkin_date": checkin_date,
            "checkout_date": checkout_date,
            "filter_by_currency": "INR",
            "longitude": long,
            "latitude": lat,
            "adults_number": num_adults,
            "order_by": "popularity",
            "units": "metric",
            "page_number": "0",
        }
        if num_children > 0:
            querystring["children_number"] = num_children
        headers = {
            "X-RapidAPI-Key": RAPID_API_KEY,
            "X-RapidAPI-Host": RAPID_API_HOST
        }

        # Send request to Booking.com API
        response = requests.get(url, headers=headers, params=querystring)
        if response.status_code == 200:
            data = response.json()
            if "result" in data:
                city_dict[city] = []
                st.write(f"Finding hotels for {city}")
                # st.write("Hotel search results for", city, ":")
                for index, hotel in enumerate(data["result"], start=1):
                    hotel_dict = {}
                    hotel_dict['hotel_name'] = hotel.get("hotel_name", "N/A")
                    price = hotel.get("min_total_price", "N/A")
                    if price is not None:
                        hotel_dict['price'] = str(price)
                    else:
                        hotel_dict['price'] = "N/A"
                    hotel_dict['address'] = hotel.get("address", "N/A")
                    hotel_dict['rating'] = hotel.get("review_score", "N/A")
                    # st.write(f"{index}. {hotel_dict['hotel_name']} - Address: {hotel_dict['address']}, Price for one day: {hotel_dict['price']} INR, Rating: {hotel_dict['rating']}")
                    city_dict[city].append(hotel_dict)
            else:
                st.write(f"No hotel results found for {city}.")
        else:
            st.write(f"Failed to retrieve hotel search results for {city}.")
    else:
        st.write(f"Invalid city name: {city}")

    return city_dict


def flight_search(input_dict):
    num_adults = input_dict['num_adults']
    cities = input_dict['cities']
    dates = input_dict['dates']
    flights = {}

    # Convert city names to IATA codes
    iata_codes = input_dict['iata_codes']

    for i in range(len(cities) - 1):
        src = cities[i]
        dest = cities[i + 1]
        flights[src] = []
        src_iata = iata_codes[i]
        dest_iata = iata_codes[i + 1]

        try:

            response = amadeus.shopping.flight_offers_search.get(
                originLocationCode=src_iata,
                destinationLocationCode=dest_iata,
                departureDate=dates[i],
                adults=min(num_adults, 9),
                travelClass='ECONOMY',
                currencyCode='INR',
                max=20
            )

            flights[src].append(response.result['data'])

        except ResponseError as error:
            print(f"Error: List of flights for {src} to {dest} is not available.")
            print(error.description())

    return flights


def display_flight_info(flight_data):
    city_flight_info = {}  # Dictionary to store flight info by city

    for src, flights in flight_data.items():
        city_flight_info[src] = []  # Initialize list for each city
        for flight_list in flights:
            for flight in flight_list:
                # Extract specific information from the flight data
                departure_time = flight['itineraries'][0]['segments'][0]['departure']['at']
                arrival_time = flight['itineraries'][0]['segments'][-1]['arrival']['at']
                airline = flight['validatingAirlineCodes'][0]
                price = flight['price']['grandTotal']
                currency = flight['price']['currency']

                # Format flight information as a dictionary
                flight_info = {
                    "Airline": airline,
                    "Departure Time": departure_time,
                    "Arrival Time": arrival_time,
                    "Price": price,
                    "Currency": currency
                }
                # Append flight information to the list for the corresponding city
                city_flight_info[src].append(flight_info)

    return city_flight_info


@st.cache_data(show_spinner=False)
def generate_itinerary(input_dict):
    # Part 1: generate the list of cities and get the hotels
    # Call the OpenAI API for creating the list of cities and dates
    input_dict['end_date'] = str(input_dict['start_date'] + timedelta(days=input_dict['num_days']))
    user_prompt = f"Generate a list of cities for a tour of {input_dict['dest']} for {input_dict['num_tourists']} " \
                  f"people with {input_dict['num_adults']} adults, purpose as {input_dict['genre']} " \
                  f"for {input_dict['num_days']} days and a budget per person of {input_dict['price_per_person']} INR starting on " \
                  f"{input_dict['start_date']}, ending on {input_dict['end_date']}. Call the function 'get_flight_hotel_info'"

    completion = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": user_prompt}],
        # Add function calling
        functions=function_descriptions,
        function_call="auto",  # specify the function call
        max_tokens=200
    )
    output = completion.choices[0].message
    cities = json.loads(output.function_call.arguments).get("loc_list")
    dates = json.loads(output.function_call.arguments).get("date_list")
    iata_codes = json.loads(output.function_call.arguments).get("iata_list")
    dates.append(input_dict['end_date'])

    print(cities)
    print(dates)
    print(iata_codes)

    input_dict['cities'] = cities
    input_dict['dates'] = dates
    input_dict['iata_codes'] = iata_codes
    all_city_dict = {}

    printables = {}
    city_string = ''
    for city in cities:
        city_string += city + ' - '
    st.subheader("Cities: ")
    st.write(city_string)
    printables['city_string'] = city_string

    for i in range(len(cities)):
        # st.write(cities[i], dates[i], dates[i+1], input_dict['num_adults'], input_dict['num_children'])
        all_city_dict.update(
           get_hotel_data(cities[i], dates[i], dates[i + 1], input_dict['num_adults'], input_dict['num_children']))
    input_dict['hotels_by_city'] = all_city_dict

    # # Part 2: Actually generate the itinerary
    # user_message = f"Design a detailed itinerary for a trip from {input_dict['src']} to {input_dict['dest']} starting from {input_dict['start_date']} and for " \
    #                f"{input_dict['num_days']} days. The ordered list of cities is {cities} and of dates is {dates}. The budget for this trip is {input_dict['price_per_person']} INR per person. This trip is designed " \
    #                f"for {input_dict['num_tourists']} mainly with their {input_dict['type_of_travelers']} with an average age of {input_dict['average_age']}.The " \
    #                f"primary interests for activities are {input_dict['genre']}.The preferred mode(s) of travel include " \
    #                f"{input_dict['mode_of_travel']}.The group prefers {input_dict['food']} food. Please structure the itinerary with a detailed " \
    #                f"plan for each day along with a every day title and no word should get repeated in a title of other days, including activities, locations, weather according to the season they are " \
    #                f"travelling and estimated travel distances and times(Do not give null values if you cannot extract information). Write the travel time and distance in the day's subheading. " \
    #                f"Ensure to consider the preferences and " \
    #                f"interests of the group for each day's schedule. Important considerations: Factor in travel time " \
    #                f"between destinations. Suggest local transportation options. Include a mix of activities that cater" \
    #                f" to the group's interests. Also add distance of travel for each day and approx time(do not give null value sif not available) " \
    #                f"of travel. Also you can give a name for each day in the itinerary which will be more " \
    #                f"appealing. Keep the response descriptive and . Give a title to the itinerary(without including the word Title) but make sure you don't repeat location names in multiple days also you can mention prime locations in title that are going to be there in iternary. Do not suggest any activities " \
    #                f"in the first city if the travel time and distance is more otherwise we can suggest activities." \
    #                f"Finally the description for each day which should look like if a human is speaking(this paragraph will be under the heading for each day)" \
    #                f"Strictly follow the number of days. Generate an itinerary for {input_dict['num_days']} days." \
                   
    user_message = f"Design a detailed {input_dict['num_days']}-day itinerary from {input_dict['src']} to {input_dict['dest']}, starting on {input_dict['start_date']}. " \
                   f"The ordered list of cities is {cities} and of dates is {dates}.The group consists of {input_dict['num_tourists']} {input_dict['type_of_travelers']} with an average age of {input_dict['average_age']}." \
                   f"Their primary interests are {input_dict['genre']} activities.The trip budget is {input_dict['price_per_person']} INR per person " \
                   f"Travel will be by {input_dict['mode_of_travel']}, and local transportation options will be included." \
                   f"Plan in detail for each day along with a every day title and no word should get repeated in a title of other days, including activities, group interests, locations, weather according to the season they are " \
                   f"travelling and estimated travel distances and times(Do not give null values if you cannot extract information).Write the travel time and distance in the day's subheading. " \
                   f"Also add distance of travel for each day and approx time(do not give null values if not available) of travel. " \
                   f"Also you can give a name for each day in the itinerary which will be more " \
                   f"appealing. Keep the response very descriptive explaining the plan for whole day(morning, afternoon, evening or anything according to you).Give a good and attractive short title to the whole itinerary(without including the word Title).Explain in such a way that a person in that group is explaing the planto other members. Do not suggest any activities " \
                   f"in the first city if the travel time and distance is more otherwise we can suggest activities." \
                   f"Also consider {special_note} if given as input from the user."\
                   f"Strictly follow the number of days and the itinerary language should be like if a travel agent is speaking to a person." \

   # Generate the travel itinerary using the modified user message
    chat_completion = client.chat.completions.create(
        messages=[
            {
                "role": "user",
                "content": user_message,
            }
        ],
        model="gpt-3.5-turbo",
        stream=True,
    )

    st.subheader("Itinerary")
    response = st.write_stream(chat_completion)

    flight_data = flight_search(input_dict)

    # Display flight information
    flight_info = display_flight_info(flight_data)
    content = response

    # Split content into individual days
    days = content.split("\n\n")

    # print(content)

    # print(response)
    locations = extract_attractive_locations(response)
    print(locations)
    delete_image_files('images')
    # response = "Dummy response"
    for line in locations.split('\n'):
        # Splitting line into day number and location name
        split_line = line.split(': ', 1)
        if len(split_line) == 2:
            day_number, location_name = split_line
            # Calling function to fetch image for location
            fetch_image(day_number, location_name, response)
        else:
            print(f"Invalid line format: {line}. Skipping.")

    banner = (input_dict['dest'])
    print("this is banner", banner)
    fetch_and_save_banner_image(banner)
    # delete_image_files("images")
    # print(titles_and_days)  # Debugging: print titles_and_days to see its structure

    st.session_state['input_dict'] = input_dict
    # with open('input_dict.pickle', 'wb') as handle:
    #     pickle.dump(input_dict, handle, protocol=pickle.HIGHEST_PROTOCOL)
    return response, all_city_dict, flight_info, days, city_string


def extract_attractive_locations(response):
    # Define the GPT-3.5 prompt
    prompt = f"Extract any and only one attractive location from the following itinerary for all days right from day 1 to the last day" \
             f"not the complete line and organize them day wise such that they will be fetched individually from a particular day" \
             f"but printed in format like day number: location name one at a time" \
             f"if there are more than one location for a day then repeat the day number" \
             f"every location must be initialized with day number and write location name and every location should be there on a new line" \
             f"here it is not needed to have city names only the attractive locations are enough\n\n{response}\n\n " \
             f"if there is no location for a particular day then simply generate any relevant location to it or about its cuisine which is not there in the data" \
             f"No location will be same, if a particular day has a location mentioned then that location will not be present in any other day" \
             f"if there are no locations for a particular day provide any travel related word that will generate an attractive image and will suit for any travel destination" \
             f"there has to be one image for everyday compulsory" \
 \
        # Generate the travel itinerary using the modified user message
    chat_completion = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {
                "role": "user",
                "content": prompt,
            }
        ],
    )

    # Get the completion text
    completion_text = chat_completion.choices[0].message.content

    return completion_text


def fetch_image(day_number, location_name, response, width=600, height=400):
    # Pexels API key (replace 'YOUR_API_KEY' with your actual Pexels API key)
    api_key = 'aX1oVcA9l4t1zj7k221MvHWgxVYZME44eCKo3szkQj3cqGqMIbyRpgdL'
    # api_key = PEXELS_API_KEY
    headers = {'Authorization': api_key}

    # Search query for location name
    query = location_name

    # Pexels API endpoint for photo search
    url = f'https://api.pexels.com/v1/search?query={query}&per_page=1'

    # Making GET request to Pexels API
    response_image = requests.get(url, headers=headers)
    data = response_image.json()

    # Check if response contains results
    if 'photos' in data and len(data['photos']) > 0:
        image_url = data['photos'][0]['src']['large']

        # Downloading image
        image_data = requests.get(image_url).content

        # Open image using Pillow
        image = PIL.Image.open(BytesIO(image_data))

        # Resize image to desired dimensions
        image = image.resize((width, height))

        # Create directory if it doesn't exist
        directory = 'images'
        if not os.path.exists(directory):
            os.makedirs(directory)

        # Saving image with filename in format 'day_number_location.jpg' inside 'images' directory
        filename = f'{directory}/{day_number}.png'
        image.save(filename, 'PNG')

        print(f"Image for {location_name} saved as {filename}")
    else:
        print(f"No image found for {location_name}")
        prompt_for_new_location(day_number, location_name, response)


def fetch_image_new_location(day_number, location_name, response, width=600, height=400):
    # Pexels API key (replace 'YOUR_API_KEY' with your actual Pexels API key)
    pexels_api_key = 'aX1oVcA9l4t1zj7k221MvHWgxVYZME44eCKo3szkQj3cqGqMIbyRpgdL'
    headers = {'Authorization': pexels_api_key}

    # Search query for location name
    query = location_name

    # Pexels API endpoint for photo search
    pexels_url = f'https://api.pexels.com/v1/search?query={query}&per_page=1'

    # Making GET request to Pexels API
    pexels_response = requests.get(pexels_url, headers=headers)
    pexels_data = pexels_response.json()

    # Check if response contains results
    if 'photos' in pexels_data and len(pexels_data['photos']) > 0:
        image_url = pexels_data['photos'][0]['src']['large']

        # Downloading image
        image_data = requests.get(image_url).content

        # Open image using Pillow
        image = PIL.Image.open(BytesIO(image_data))

        # Resize image to desired dimensions
        image = image.resize((width, height))

        # Create directory if it doesn't exist
        directory = 'images'
        if not os.path.exists(directory):
            os.makedirs(directory)

        # Saving image with filename in format 'day_number_location.jpg' inside 'images' directory
        filename = f'{directory}/Day {day_number}.png'
        image.save(filename, 'PNG')

        print(f"Image for {location_name} saved as {filename}")
    else:
        # If no image found on Pexels, call function to generate another attractive location
        # print(f"No image found for {location_name}. Generating another attractive location for day {day_number}...")
        prompt_for_new_location(day_number, location_name, response)


def prompt_for_new_location(day_number, location_name, response):
    prompt = f"there is no use for {location_name} so extract another new location for{day_number}from {response}" \
             f"the name should not be same as the previous location it should be new and can be cuisine or anything that is attractive to see" \
             f"provide only the location name in format Day Day number: new location name, do not provide whole response" \
        # Generate the travel itinerary using the modified user message
    new_location = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {
                "role": "user",
                "content": prompt,
            }
        ],
    )

    # Get the completion text
    completion_text = new_location.choices[0].message.content
    print(completion_text)
    # print(f"Generated new attractive location for day {day_number}: {completion_text}")
    fetch_image_new_location(day_number, completion_text)
    return completion_text
    # Call fetch_image with the new location


def fetch_and_save_banner_image(banner, width=600, height=400):
    # Pexels API key (replace 'YOUR_API_KEY' with your actual Pexels API key)
    api_key = PEXELS_API_KEY
    headers = {'Authorization': api_key}

    # Search query for location name
    query = banner

    # Pexels API endpoint for photo search
    url = f'https://api.pexels.com/v1/search?query={query}&per_page=1'

    # Making GET request to Pexels API
    response = requests.get(url, headers=headers)
    data = response.json()

    # Check if response contains results
    if 'photos' in data and len(data['photos']) > 0:
        image_url = data['photos'][0]['src']['large']

        # Downloading image
        image_data = requests.get(image_url).content

        # Open image using Pillow
        image = PIL.Image.open(BytesIO(image_data))

        image = image.resize((width, height))
        # Create directory if it doesn't exist
        directory = 'images'
        if not os.path.exists(directory):
            os.makedirs(directory)

        # Saving image with filename in format 'day_number_location.jpg' inside 'images' directory
        filename = f'{directory}/banner.png'
        image.save(filename, 'PNG')

        print(f"Image for {banner} saved as {filename}")
    else:
        print(f"No image found for {banner}")


def delete_image_files(directory):
    # Get a list of all files in the directory
    files = os.listdir(directory)

    # Loop through each file
    for file in files:
        # Check if the file is an image file (you can customize this condition as needed)
        if file.endswith(".jpg") or file.endswith(".jpeg") or file.endswith(".png") or file.endswith(".gif"):
            # Construct the full path to the file
            file_path = os.path.join(directory, file)
            try:
                # Attempt to remove the file
                os.remove(file_path)
                print(f"Deleted: {file_path}")
            except Exception as e:
                print(f"Error deleting {file_path}: {e}")


def remove_text(doc, text):
    #   text="Evaluation Warning: The document was created with Spire.Doc for Python"
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if text in run.text:
                run.text = run.text.replace(text, "")


def extract_day_number(filename):
    # If filename is 'cover_page.docx', return a very low number to prioritize it
    if filename.lower() == 'cover_page.docx':
        return float('-inf')
    else:
        # Split filename by underscores
        parts = filename.split('_')
        # Check if the filename has at least two parts
        if len(parts) >= 2:
            # Extract the second part
            day_part = parts[1]
            # Try to convert the day part to an integer
            try:
                day_number = int(day_part)
            except ValueError:
                # If conversion fails, return a large number
                return float('inf')
            else:
                return day_number
        else:
            # If the filename does not have at least two parts, return a large number
            return float('inf')


# Function to extract the day number from the filename
def extract_day_number(filename):
    match = re.search(r'day_(\d+)_', filename)
    if match:
        return int(match.group(1))
    elif 'cover_page' in filename:
        return 0  # Assigning 0 for cover page to ensure it comes first
    else:
        return float('inf')


def text_to_doc(itinerary, input_dict):
    day_itineraries = generate_day_itineraries(itinerary)
    city_names = ", ".join(input_dict['cities'])
    folder_name = "generated_itineraries"
    fetch_and_save_banner_image(itinerary.split('\n')[0])
    # Create the directory if it doesn't exist
    if not os.path.exists(folder_name):
        os.makedirs(folder_name)
       

    # Delete the previously generated documents
    for filename in os.listdir(folder_name):
        file_path = os.path.join(folder_name, filename)
        try:
            if os.path.isfile(file_path):
                os.remove(file_path)
        except Exception as e:
            print(f"Error deleting {file_path}: {e}")

    first_page = DocxTemplate('mergeDocs/page1.docx')

    # image_path = 'images/banner.png'
    context = {
        'tour_heading': itinerary.split('\n')[0],
        'num_days': input_dict['num_days'],
        'budget': input_dict['price_per_person'],
        'cities': city_names,
        # 'day_image': InlineImage(
        #     first_page, image_path, width=Mm(70), height=Mm(70))
    }
    first_page.render(context)
    # first_page.replace_media('mergeDocs/front_img.png', 'images/banner.png')

    file_path = os.path.join(folder_name, 'cover_page.docx')
    first_page.save(file_path)

    # Load the template document
    tpl = DocxTemplate("mergeDocs/page2.docx")

    for day_number, day_itinerary in day_itineraries.items():
        # Extract the first line of the itinerary
        first_line = day_itinerary.split('\n')[0]
        # print("Inside the text_to_doc func: ", day_itinerary, 'First line: ', first_line, 'Second line: ',
        #       day_itinerary)
        # Join city names into a comma-separated string
        first_newline_index = day_itinerary.find('\n')

        # Check if '\n' exists in the string
        if first_newline_index != -1:
            # Extract the substring starting from the index after the first '\n'
            day_itinerary = day_itinerary[first_newline_index + 1:]

        # Extract the image file path based on the day_number
        image_folder = "images"
        image_file = f"Day {day_number}.png"
        image_path = os.path.join(image_folder, image_file)

        # Define the context dictionary
        context = {
            'tour_heading': first_line,
            'num_days': input_dict['num_days'],
            'budget': input_dict['price_per_person'],
            'day_itinerary': day_itinerary,
            'day_title': first_line,
            'day_image': InlineImage(
                tpl, image_path, width=Mm(100), height=Mm(80))
        }

        # Replace placeholders in the document
        tpl.render(context)

        os.makedirs(folder_name, exist_ok=True)

        # Modify the file path where the documents are saved
        file_path = os.path.join(folder_name, f'day_{day_number}_itinerary.docx')
        tpl.save(file_path)

    # Create a Document object
    destDoc = Document()

    # Define the folder path containing the files to merge
    folder_path = "generated_itineraries"

    # Create the directory if it doesn't exist
    os.makedirs(folder_path, exist_ok=True)

    # List all files in the folder
    files_to_merge = os.listdir(folder_path)

    # Filter only the .docx files
    files_to_merge = [file for file in files_to_merge if file.endswith('.docx')]
    files_to_merge.sort(key=extract_day_number)
    cover_page = None
    # Loop through the list
    for file in files_to_merge:
        # Construct the full file path
        file_path = os.path.join(folder_path, file)

        # Load the source document
        sourceDoc = Document()
        sourceDoc.LoadFromFile(file_path)

        # Add the cover page at the beginning if found

        # Keep the formatting of the source document when it is merged
        # sourceDoc.KeepSameFormat = True

        # Import the content from the document into the destination document
        destDoc.ImportContent(sourceDoc)

    # Save the result document
    destDoc.SaveToFile("Itinerary1.docx", FileFormat.Docx2016)
    doc = docx.Document("Itinerary1.docx")
    remove_text(doc, "Evaluation Warning: The document was created with Spire.Doc for Python")
    doc.save("Itinerary.docx")
    destDoc.Close()
    sourceDoc.Close()

    with open("Itinerary.docx", "rb") as file:
        bytes_content = file.read()
    return bytes_content

def enter_details():
    input_dict_hotel = {}
    input_dict_car = {}
    num_entries = st.number_input("Enter the number of entries:", min_value=1, step=1)

    for i in range(num_entries):
        city = st.text_input(f"City {i+1}")
        hotel = st.text_input(f"Hotel for City {i+1}")
        price = st.number_input(f"Price per Night for Hotel in City {i+1}")
        car = st.text_input(f"Car for City {i+1}")
        fare = st.number_input(f"Fare for Car in City {i+1}")
        input_dict_hotel[city] = (hotel, price)
        input_dict_car[city] = (car, fare)
    
    return input_dict_hotel, input_dict_car


def get_day_itinerary(days, day_number):

    for day in days:
        # print(day)
        if day.startswith(str(day_number) + ":"):
            day = day.replace('*', '')
            day = day.replace('###', '')
            return "Day " + day


def generate_day_itineraries(itinerary):
    day_itineraries = {}

    # Regular expression pattern to extract day numbers
    pattern = r"Day (\d+)"

    # Find all matches of the pattern in the itinerary
    matches = re.findall(pattern, itinerary)

    # Convert matched day numbers to integers and get the maximum day
    max_day = max(int(day) for day in matches)

    # Initialize day itineraries for all days up to the maximum day
    for day_number in range(1, max_day + 1):
        day_itineraries[day_number] = ""

    # Generate day itineraries for each day
    def merge_strings_with_previous(lst):
        merged_list = [""]
        day_idx = 0
        for item in lst:
            if not item[0].isdigit():
                tail = "Day " + item
                merged_list[day_idx] += tail
            else:
                day_idx += 1
                merged_list.append(item)

        return merged_list

    # Split the itinerary into days
    days = itinerary.split("Day ")

    days = merge_strings_with_previous(days)
    for day_number in range(1, max_day + 1):
        day_itinerary = get_day_itinerary(days, day_number)
        if day_itinerary:
            day_itineraries[day_number] = day_itinerary
    return day_itineraries


if st.session_state.get('input_dict', False):
    for key in input_dict.keys():
        if input_dict[key] != st.session_state['input_dict'][key]:
            st.session_state['data_changed'] = True
            break

if st.button("Generate Itinerary", type="primary"):
    # null_flag = False
    # if 'OPENAI_API_KEY' not in env_vars or 'X-RapidAPI-Key' not in env_vars or 'AMADEUS_API_KEY' not in env_vars or 'AMADEUS_API_SECRET' not in env_vars or 'PEXELS_API_KEY' not in env_vars:
    #     st.warning('Enter all the API keys')
    #     null_flag = True
    # for key in input_dict.keys():
    #     if input_dict[key] is None:
    #         st.warning(f'Please enter {key}!')
    #         null_flag = True
    #         break

    generated_itinerary, city_dict, flight_info, days, city_string = generate_itinerary(input_dict)
    st.session_state["cached_data_generated"] = True
    st.session_state['data_changed'] = False
    isGenerated = True

elif st.session_state.get("cached_data_generated", False) and not st.session_state['data_changed']:
    generated_itinerary, city_dict, flight_info, days, city_string = generate_itinerary(input_dict)

if st.session_state.get("cached_data_generated", False) and not st.session_state['data_changed']:
    st.subheader("Hotels")
    for city, hotels in city_dict.items():
        city_expander = st.expander(f"{city}")
        with city_expander:
            for hotel in hotels:
                price = round(float(hotel['price']))
                num_tourists = input_dict['num_adults'] + input_dict['num_children']
                price_per_person = price // num_tourists

                st.write(f"- {hotel['hotel_name']}")
                st.write(f"  Address: {hotel['address']}")
                st.write(f"  Price per day: ₹ {price_per_person} ")
                st.write(f"  Rating: {hotel['rating']}")
                # Add more details as needed (amenities, images, etc.)
                st.write("---")  # Separator between hotels\

    st.subheader("Flight Details")
    for city, flights in flight_info.items():
        city_expander = st.expander(f"{city}")
        with city_expander:
            for flight in flights:
                st.write(f"- {flight['Airline']}")
                st.write(f"  Departure Time: {flight['Departure Time']}")
                st.write(f"  Arrival Time: {flight['Arrival Time']}")
                st.write(f"  Price: ₹ {flight['Price']}")
                # Add more details as needed (amenities, images, etc.)
                st.write("---")

    text_to_doc(generated_itinerary, st.session_state['input_dict'])

    input_dict_hotel, input_dict_car = enter_details()
    document = generate_word_document(input_dict_hotel, input_dict_car)
    document.save('dynamic_tables.docx')

    master = Document_compose('Itinerary.docx')
    composer = Composer(master)
    doc2 = Document_compose('dynamic_tables.docx')
    composer.append(doc2)
    # Append doc3 (Accomodation)
    doc3 = Document_compose('lastpage.docx')  # Provide the path to the third document
    composer.append(doc3)
    composer.save("combined.docx")

    # st.download_button(
    #     label="Download Word Document",
    #     data='combined.docx',
    #     file_name=f"{input_dict['dest']} Itinerary.docx",
    #     mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",

    # )
    # Convert the combined.docx file to bytes
    with open("combined.docx", "rb") as file:
        doc_bytes = file.read()

    # Create a BytesIO object from the bytes
    doc_io = io.BytesIO(doc_bytes)

    # Display the download button with the BytesIO object
    st.download_button(
        label="Download Word Document",
        data=doc_io,
        file_name=f"{input_dict['dest']} Itinerary.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )