from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docxcompose.composer import Composer
from docx import Document as Document_compose

def generate_word_document(input_dict_hotel, input_dict_car):
    # Create a new Document object
    document = Document()

    # Add a page break
    document.add_page_break()

    # Add a paragraph for the table heading
    paragraph=document.add_paragraph("Hotel Details", style='Heading 1')
    run = paragraph.runs[0]
    run.font.size = Pt(18)

    # Define table data for the first table (Hotel) with column headings
    table_data1 = [["Destination", "Hotel", "Price per Night"]]

    # Add data to table_data1
    for city, details in input_dict_hotel.items():
        hotel, price = details
        table_data1.append([city, hotel, price])

    # Add the first table (Hotel) to the document
    table1 = document.add_table(rows=len(table_data1), cols=3)

    # Adding data to the first table (Hotel)
    for row_data in table_data1:
        row = table1.add_row().cells
        for i, cell_data in enumerate(row_data):
            cell = row[i]
            paragraph = cell.paragraphs[0]
            paragraph.text = str(cell_data)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(14)
                run.font.name = 'Arial'
                run.font.color.rgb = RGBColor(0, 0, 0)

    # Apply custom table style to the first table (Hotel)
    # table1.style = 'Colorful Grid Accent 5'
    table1.style = 'Colorful Grid Accent 2'

    # Add a paragraph to separate tables
    document.add_paragraph()

    # Add a paragraph for the table heading for the second table (Car)
    paragraph=document.add_paragraph("Car Details", style='Heading 1')
    run = paragraph.runs[0]
    run.font.size = Pt(18)

    # Define table data for the second table (Car) with column headings
    table_data2 = [["Destination", "Car", "Fare"]]

    # Add data to table_data2
    for city, details in input_dict_car.items():
        car, fare = details
        table_data2.append([city, car, fare])

    # Add the second table (Car) to the document
    table2 = document.add_table(rows=len(table_data2), cols=3)

    # Adding data to the second table (Car)
    for row_data in table_data2:
        row = table2.add_row().cells
        for i, cell_data in enumerate(row_data):
            cell = row[i]
            paragraph = cell.paragraphs[0]
            paragraph.text = str(cell_data)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(14)
                run.font.name = 'Arial'
                run.font.color.rgb = RGBColor(0, 0, 0)

    # Apply custom table style to the second table (Car)
    table2.style = 'Colorful Grid Accent 2'

    return document

def add_background_image(document, image_path):
    section = document.sections[0]  # Get the first section of the document
    header = section.header
    header.is_linked_to_previous = False  # Ensure the header is not linked to the previous section
    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()  # Add a paragraph if not exists
    run = paragraph.add_run()
    run.add_picture(image_path, width=Inches(8.5), height=Inches(11))  # Adjust width and height as needed
    return document

# def enter_details():
#     input_dict_hotel = {}
#     input_dict_car = {}
#     num_entries = st.number_input("Enter the number of entries:", min_value=1, step=1)

#     for i in range(num_entries):
#         city = st.text_input(f"City {i+1}")
#         hotel = st.text_input(f"Hotel for City {i+1}")
#         price = st.number_input(f"Price per Night for Hotel in City {i+1}")
#         car = st.text_input(f"Car for City {i+1}")
#         fare = st.number_input(f"Fare for Car in City {i+1}")
#         input_dict_hotel[city] = (hotel, price)
#         input_dict_car[city] = (car, fare)
    
#     return input_dict_hotel, input_dict_car

# input_dict_hotel, input_dict_car = enter_details()
# document = generate_word_document(input_dict_hotel, input_dict_car)
# document.save('dynamic_tables.docx')

# master = Document_compose('Itinerary.docx')
# composer = Composer(master)
# doc2 = Document_compose('dynamic_tables.docx')
# composer.append(doc2)
# composer.save("combined.docx")
